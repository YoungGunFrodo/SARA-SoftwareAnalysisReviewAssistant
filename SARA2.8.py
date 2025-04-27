import sys
import os
import re
import base64
import requests
import fitz  # PyMuPDF
import vt
from bs4 import BeautifulSoup
from duckduckgo_search import DDGS
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QLabel, QLineEdit, QPushButton,
    QProgressBar, QTextEdit, QTabWidget, QCheckBox, QPlainTextEdit, QFileDialog, QHBoxLayout,
    QTableWidget, QTableWidgetItem
)
from PySide6.QtCore import Qt, QThread, Signal
import qdarkstyle
from docx import Document
from docx.shared import Pt
import io
from datetime import datetime
import json
import traceback


# Load config.json
with open("config.json", "r") as f:
    config = json.load(f)

# Use values from config
VIRUSTOTAL_API_KEY = config["VIRUSTOTAL_API_KEY"]
ABUSEIPDB_API_KEY = config["ABUSEIPDB_API_KEY"]
OLLAMA_MODEL = config["OLLAMA_MODEL"]
CHATGPT_API_KEY = config["CHATGPT_API_KEY"]
TENANT_ID = config["TENANT_ID"]
CLIENT_ID = config["CLIENT_ID"]
CLIENT_SECRET = config["CLIENT_SECRET"]
DRIVE_ID = config["DRIVE_ID"]
EXCEL_FILE_ID = config["EXCEL_FILE_ID"]
WORKSHEET_NAME = config["WORKSHEET_NAME"]
CHATGPT_MODEL = "gpt-4.1-mini"

# === Upload ===
class SharePointUploader:
    def __init__(self, tenant_id, client_id, client_secret, drive_id):
        self.tenant_id = tenant_id
        self.client_id = client_id
        self.client_secret = client_secret
        self.drive_id = drive_id
        self.access_token = self.get_access_token()

    def get_access_token(self):
        url = f"https://login.microsoftonline.com/{self.tenant_id}/oauth2/v2.0/token"
        data = {
            "grant_type": "client_credentials",
            "client_id": self.client_id,
            "client_secret": self.client_secret,
            "scope": "https://graph.microsoft.com/.default"
        }
        return requests.post(url, data=data).json()["access_token"]

    def upload_file(self, buffer, sharepoint_filename="report.docx", folder="General"):
     headers = {
         "Authorization": f"Bearer {self.access_token}",
         "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
     }
     url = f"https://graph.microsoft.com/v1.0/drives/{self.drive_id}/root:/{folder}/{sharepoint_filename}:/content"
     response = requests.put(url, headers=headers, data=buffer.read())

     # Optional debug: print the status code and message
     print("Upload status:", response.status_code)
     print("Response:", response.text)

     return response.status_code in [200, 201]

# === LLM Workers ===
class LocalLLMWorker(QThread):
    result = Signal(str)
    prompt_output = Signal(str)

    def __init__(self, prompt):
        super().__init__()
        self.prompt = prompt

    def run(self):
        self.prompt_output.emit(self.prompt)
        try:
            response = requests.post("http://localhost:11434/api/generate", json={
                "model": OLLAMA_MODEL,
                "prompt": self.prompt,
                "stream": False
            })
            result = response.json().get("response", "LLM response unavailable")
        except Exception as e:
            result = f"LLM error: {e}"
        self.result.emit(result)

class ChatGPTWorker(QThread):
    result = Signal(str)
    prompt_output = Signal(str)

    def __init__(self, prompt):
        super().__init__()
        self.prompt = prompt

    def run(self):
        try:
            # Show the prompt in your GUI (if the checkbox is enabled)
            self.prompt_output.emit(self.prompt)

            # Build headers & payload
            headers = {
                "Authorization": f"Bearer {CHATGPT_API_KEY}",
                "Content-Type": "application/json"
            }
            data = {
                "model": CHATGPT_MODEL,          # pulled from config.json
                "messages": [
                    {"role": "system", "content": "You are a helpful cybersecurity assistant."},
                    {"role": "user", "content": self.prompt}
                ]
            }

            # Fire the request with a timeout so it never hangs forever
            response = requests.post(
                "https://api.openai.com/v1/chat/completions",
                headers=headers,
                json=data,
                timeout=30
            )

            # HTTP‐level errors
            if response.status_code != 200:
                raise RuntimeError(f"API error {response.status_code}: {response.text}")

            # Parse JSON
            body = response.json()
            if "choices" not in body or not body["choices"]:
                raise RuntimeError("Unexpected response format from ChatGPT API")

            # Success!
            result = body["choices"][0]["message"]["content"]

        except Exception as e:
            # Capture the full traceback
            tb = traceback.format_exc()
            result = f"ChatGPT error: {e}\n\n{tb}"

            # Persist the error in your app’s own log folder
            try:
                log_dir = os.path.join(os.path.expanduser("~"), ".sara_logs")
                os.makedirs(log_dir, exist_ok=True)
                with open(os.path.join(log_dir, "chatgpt_errors.txt"), "a", encoding="utf-8") as f:
                    f.write(f"{datetime.now().isoformat()} — {result}\n\n")
            except Exception:
                pass  # if logging fails, we still want to emit the error back to the UI

        # Always emit *some* result, so the GUI thread wakes up
        self.result.emit(result)

# === VirusTotal ===
class VirusTotalWorker(QThread):
    result = Signal(str)

    def __init__(self, url):
        super().__init__()
        self.url = url

    def run(self):
     try:
         with vt.Client(VIRUSTOTAL_API_KEY) as client:
             url_id = base64.urlsafe_b64encode(self.url.encode()).decode().strip("=")
             result = client.get_object(f"/urls/{url_id}")
             details = [
                 f"URL: {self.url}",
                 f"Harmless: {result.last_analysis_stats['harmless']}",
                 f"Malicious: {result.last_analysis_stats['malicious']}",
                 f"Suspicious: {result.last_analysis_stats['suspicious']}",
                 f"Undetected: {result.last_analysis_stats['undetected']}",
                 f"Scan ID: {result.id}"
             ]
             self.result.emit("\n".join(details))
     except Exception as e:
         self.result.emit(f"VirusTotal error: {e}")

# === AbuseIPDB ===
class AbuseIPDBWorker(QThread):
    result = Signal(str)

    def __init__(self, domain):
        super().__init__()
        self.domain = domain

    def run(self):
     try:
         ip_resp = requests.get(f"https://dns.google/resolve?name={self.domain}&type=A", timeout=10)
         ip_data = ip_resp.json()
         ip = ip_data.get("Answer", [{}])[0].get("data", "")
         if not ip:
             self.result.emit("AbuseIPDB error: Could not resolve IP")
             return
         headers = {"Key": ABUSEIPDB_API_KEY, "Accept": "application/json"}
         resp = requests.get(f"https://api.abuseipdb.com/api/v2/check?ipAddress={ip}&maxAgeInDays=90", headers=headers)
         d = resp.json().get("data", {})
         details = [
             f"Checked IP: {d.get('ipAddress')}",
             f"Abuse Score: {d.get('abuseConfidenceScore')}",
             f"Reports: {d.get('totalReports')}",
             f"Country: {d.get('countryCode')}",
             f"ISP: {d.get('isp')}",
             f"Domain: {d.get('domain')}"
         ]
         self.result.emit("\n".join(details))
     except Exception as e:
         self.result.emit(f"AbuseIPDB error: {e}")

# === Microsoft Graph API Excel Integration ===
class ExcelConnector:
    def __init__(self, tenant_id, client_id, client_secret, file_id):
        self.token = self._get_token(tenant_id, client_id, client_secret)
        self.file_id = file_id

    def _get_token(self, tenant_id, client_id, client_secret):
        url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"
        data = {
            "grant_type": "client_credentials",
            "client_id": client_id,
            "client_secret": client_secret,
            "scope": "https://graph.microsoft.com/.default"
        }
        return requests.post(url, data=data).json().get("access_token")

    def get_incomplete_requests(self):
      headers = {"Authorization": f"Bearer {self.token}"}
      url = f"https://graph.microsoft.com/v1.0/drives/{DRIVE_ID}/items/{EXCEL_FILE_ID}/workbook/worksheets('{WORKSHEET_NAME}')/usedRange?valuesOnly=true"
      response = requests.get(url, headers=headers).json()
      values = response.get("values", [])
      if not values or len(values) < 2:
          return []

      headers_row = values[0]
      comment_index = headers_row.index("Reviewer Comments") if "Reviewer Comments" in headers_row else -1

      # Only include rows where the comment is empty
      data = []
      for idx, row in enumerate(values[1:], start=2):  # +2 accounts for 0-indexing and header row
         if row and row[0] and "complete" not in str(row).lower():
             data.append({"__row__": idx, **dict(zip(headers_row, row))})

      return data  

    def update_row(self, excel_row, comment):
      headers = {
           "Authorization": f"Bearer {self.token}",
          "Content-Type": "application/json"
      }

      # Fetch current values in the row to avoid wiping anything out
      get_url = f"https://graph.microsoft.com/v1.0/drives/{DRIVE_ID}/items/{EXCEL_FILE_ID}/workbook/worksheets('{WORKSHEET_NAME}')/range(address='A{excel_row}:Z{excel_row}')"
      response = requests.get(get_url, headers=headers).json()
      existing_values = response.get("values", [[]])[0]

      # Pad row if it's shorter than needed
      while len(existing_values) < 10:
          existing_values.append("")

      # Update status and comment
      existing_values[0] = "complete"
      existing_values[9] = comment  # Adjust this if your comment column moves

      # Write back to Excel
      patch_url = f"https://graph.microsoft.com/v1.0/drives/{DRIVE_ID}/items/{EXCEL_FILE_ID}/workbook/worksheets('{WORKSHEET_NAME}')/range(address='A{excel_row}:J{excel_row}')"
      patch_data = {"values": [existing_values[:10]]}
      requests.patch(patch_url, headers=headers, json=patch_data)

# === Requests Tab Widget ===
class RequestTab(QWidget):
    def __init__(self, file_id, tenant_id, client_id, client_secret):
        super().__init__()
        self.file_id = file_id
        self.tenant_id = tenant_id
        self.client_id = client_id
        self.client_secret = client_secret

        self.connector = ExcelConnector(tenant_id, client_id, client_secret, file_id)
        self.layout = QVBoxLayout()
        self.setLayout(self.layout)

        self.table = QTableWidget()
        self.table.setColumnCount(0)
        self.table.setRowCount(0)
        self.layout.addWidget(self.table)
        # Comment box + complete button
        self.comment_box = QTextEdit()
        self.comment_box.setPlaceholderText("Optional comments before marking complete...")
        self.complete_button = QPushButton("Mark Complete")

        # Wire up the button
        self.complete_button.clicked.connect(self.mark_complete)

        # Add to layout
        self.layout.addWidget(self.comment_box)
        self.layout.addWidget(self.complete_button)
        self.refresh_requests()

    def refresh_requests(self):
        data = self.connector.get_incomplete_requests()
        if not data:
            self.table.setRowCount(0)
            self.table.setColumnCount(1)
            self.table.setHorizontalHeaderLabels(["No requests found"])
            return

        self.table.setRowCount(len(data))
        self.table.setColumnCount(len(data[0]))
        self.table.setHorizontalHeaderLabels(list(data[0].keys()))

        for row_idx, row in enumerate(data):
            for col_idx, (key, value) in enumerate(row.items()):
                item = QTableWidgetItem(str(value))
                if key == "__row__":
                 item.setData(Qt.UserRole, value)  # Save actual Excel row
                self.table.setItem(row_idx, col_idx, item)

        self.table.resizeColumnsToContents()
        self.table.horizontalHeader().setStretchLastSection(True)

    def mark_complete(self):
      selected = self.table.currentRow()
      if selected == -1:
           return

       # Get the actual Excel row (stored during refresh)
      excel_row = int(self.table.item(selected, 0).data(Qt.UserRole))  # read from UserRole

      comment = self.comment_box.toPlainText().strip()
      self.connector.update_row(excel_row, comment)
      self.refresh_requests()
      self.comment_box.clear()

# === Main App ===
class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("SARA - Software Analysis & Review Assistant")
        self.setMinimumSize(1000, 700)

        self.url_input = QLineEdit()
        self.url_input.setPlaceholderText("Enter vendor website URL...")
        self.local_button = QPushButton("Review with Local LLM")
        self.gpt_button = QPushButton("Review with ChatGPT")
        self.status_label = QLabel("Status: Waiting")
        self.progress = QProgressBar(); self.progress.setRange(0, 100)
        self.export_local_button = QPushButton("Export Locally")
        self.export_teams_button = QPushButton("Export to SharePoint")

        self.tabs = QTabWidget()
        self.llm_tab, self.vt_tab, self.abuse_tab = QTextEdit(), QTextEdit(), QTextEdit()
        self.rec_tab = QWidget()
        for tab in (self.llm_tab, self.vt_tab, self.abuse_tab): tab.setReadOnly(True)

        self.rec_layout = QVBoxLayout(self.rec_tab)
        self.rec_text = QTextEdit(); self.rec_text.setReadOnly(True)
        self.rec_layout.addWidget(self.rec_text)
        self.show_prompt_chk = QCheckBox("Show full LLM prompt")
        self.prompt_box = QPlainTextEdit()
        self.prompt_box.setReadOnly(True)
        self.prompt_box.hide()
        self.show_prompt_chk.toggled.connect(lambda x: self.prompt_box.setVisible(x))
        self.rec_layout.addWidget(self.show_prompt_chk)
        self.rec_layout.addWidget(self.prompt_box)

        # Create requests tab and pass API config
        self.requests_tab = RequestTab(EXCEL_FILE_ID, TENANT_ID, CLIENT_ID, CLIENT_SECRET)
        self.tabs.addTab(self.requests_tab, "Requests")

        # Add existing tabs
        for name, tab in zip(["LLM Summary", "VirusTotal", "AbuseIPDB", "Accessed Records"],
                             [self.llm_tab, self.vt_tab, self.abuse_tab, self.rec_tab]):
         self.tabs.addTab(tab, name)


        layout = QVBoxLayout()
        layout.addWidget(self.url_input)
        review_buttons = QHBoxLayout()
        review_buttons.addWidget(self.local_button)
        review_buttons.addWidget(self.gpt_button)
        layout.addLayout(review_buttons)
        layout.addWidget(self.status_label)
        layout.addWidget(self.progress)
        export_buttons = QHBoxLayout()
        export_buttons.addWidget(self.export_local_button)
        export_buttons.addWidget(self.export_teams_button)
        layout.addLayout(export_buttons)
        layout.addWidget(self.tabs)
        container = QWidget(); container.setLayout(layout)
        self.setCentralWidget(container)

        self.local_button.clicked.connect(lambda: self.start_review("local"))
        self.gpt_button.clicked.connect(lambda: self.start_review("gpt"))
        self.export_local_button.clicked.connect(self.export_local)
        self.export_teams_button.clicked.connect(self.export_sharepoint)

        self.latest_records = ""
        self.latest_llm = ""

    def start_review(self, mode):
        url = self.url_input.text().strip()
        if not url: return
        self.status_label.setText("Scraping site...")
        self.progress.setValue(5)

        html = requests.get(url, timeout=10).text
        soup = BeautifulSoup(html, "html.parser")
        site_text = soup.get_text()[:20000]
        pdf_links = [requests.compat.urljoin(url, a['href']) for a in soup.find_all("a", href=True) if a['href'].endswith(".pdf")]
        web_summary, web_links = self.web_search(url)

        pdf_text = ""
        for link in pdf_links:
            try:
                r = requests.get(link); open("temp.pdf", "wb").write(r.content)
                doc = fitz.open("temp.pdf")
                pdf_text += f"\n--- PDF: {link} ---\n" + "\n".join([p.get_text() for p in doc])
                doc.close(); os.remove("temp.pdf")
            except: pass

        self.latest_records = f"Scanned: {url}\nWeb Sources:\n" + "\n".join(web_links) + "\nPDFs:\n" + "\n".join(pdf_links)
        self.rec_text.setPlainText(self.latest_records)

        domain = url.split("//")[-1].split("/")[0]
        self.vt_thread = VirusTotalWorker(url); self.vt_thread.result.connect(self.vt_tab.setPlainText); self.vt_thread.start()
        self.abuse_thread = AbuseIPDBWorker(domain); self.abuse_thread.result.connect(self.abuse_tab.setPlainText); self.abuse_thread.start()

        prompt = f"""
You are a cybersecurity assistant reviewing software for data and security concerns for use in a university environment. You respond clearly, are cautious, thorough, and knowledgable of cybersecurity and software analysis.
Analyze the following data from the vendor's website and the web search:

Website Text:
{site_text}

Web Search Summary:
{web_summary}

PDF Content:
{pdf_text}

Respond with:
- Security concerns
- Cloud or desktop app
- Are admin rights needed to run
- Collected data and where it's stored
- Encryption use and type of encryption
- Internet requirement
- FERPA compliance
- HIPAA compliance
- PHI compliance
- PII compliance
- Use of AI
- Primary purpose
- Software type
- Any other concerns or insights about this software in a university setting 
        """
        self.status_label.setText("Analyzing...")
        self.progress.setValue(50)

        worker = LocalLLMWorker(prompt) if mode == "local" else ChatGPTWorker(prompt)
        worker.result.connect(self.llm_tab.setPlainText)
        worker.prompt_output.connect(self.prompt_box.setPlainText)
        worker.finished.connect(lambda: self.progress.setValue(100))
        worker.finished.connect(lambda: self.status_label.setText("Status: Done"))
        worker.start()

    def web_search(self, url):
        terms = ["security vulnerabilities", "data privacy", "FERPA HIPAA"]
        results, links = [], []
        with DDGS() as ddgs:
            for t in terms:
                for r in ddgs.text(f"{url} {t}", max_results=3):
                    results.append(f"- {r['title']}: {r['body']}")
                    links.append(r['href'])
        return "\n".join(results), links

    def export_local(self):
        self.export_report(local=True)

    def export_sharepoint(self):
     from io import BytesIO
     doc = Document()
     doc.add_heading("SARA - Software Assessment Report", 0)
     for label, content in zip(["LLM Summary", "VirusTotal", "AbuseIPDB", "Accessed Records"],
                               [self.llm_tab.toPlainText(), self.vt_tab.toPlainText(), self.abuse_tab.toPlainText(), self.latest_records]):
         doc.add_heading(label, level=1)
         doc.add_paragraph(content).style.font.size = Pt(10)
          
     domain = self.url_input.text().split("//")[-1].split("/")[0]
     timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
     filename = f"{domain}_{timestamp}.docx"
      
     buffer = BytesIO()
     doc.save(buffer)
     buffer.seek(0)

     uploader = SharePointUploader(TENANT_ID, CLIENT_ID, CLIENT_SECRET, DRIVE_ID)
     success = uploader.upload_file(buffer, sharepoint_filename=filename)
     self.status_label.setText("✅ Uploaded" if success else "❌ Upload Failed")

    def export_report(self, local=True):
        doc = Document()
        doc.add_heading("SARA - Software Assessment Report", 0)
        for label, content in zip(["LLM Summary", "VirusTotal", "AbuseIPDB", "Accessed Records"],
                                  [self.llm_tab.toPlainText(), self.vt_tab.toPlainText(), self.abuse_tab.toPlainText(), self.latest_records]):
            doc.add_heading(label, level=1)
            doc.add_paragraph(content).style.font.size = Pt(10)

        domain = self.url_input.text().split("//")[-1].split("/")[0]
        filename = f"{domain}.docx"
        if not local:
            doc.save(filename)
            uploader = SharePointUploader(TENANT_ID, CLIENT_ID, CLIENT_SECRET, DRIVE_ID)
            success = uploader.upload_file(filename, sharepoint_filename=filename)
            self.status_label.setText("✅ Uploaded" if success else "❌ Upload Failed")
        else:
            f, _ = QFileDialog.getSaveFileName(self, "Save Report", filename, "Word Files (*.docx)")
            if f:
                doc.save(f)
                self.status_label.setText("✅ Saved locally")

if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setStyleSheet(qdarkstyle.load_stylesheet())
    win = MainWindow()
    win.show()
    sys.exit(app.exec())



